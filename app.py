# -*- coding:utf-8 -*-

import os, base64, time, json
from docx import Document
import tornado.ioloop
import tornado.web
from tornado.web import RequestHandler, Application
from OCRsetting import static_path, IP, PORT
from pymongo import MongoClient

from bson import json_util
from bson.objectid import ObjectId
import uuid
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfgen import canvas

# 注册字体
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib import fonts

pdfmetrics.registerFont(TTFont("SimSun", "static/fonts/SimSun.ttf"))
fonts.addMapping("SimSun", 0, 0, "SimSun")
fonts.addMapping("SimSun", 0, 1, "SimSun")

MONGODB_DB_URL = 'mongodb://localhost:27017/'
MONGODB_DB_NAME = 'ocr_db'

client = MongoClient(MONGODB_DB_URL)
db = client[MONGODB_DB_NAME]

'''
users : {
    name      : 
    password  : 
    ocr_times : 
}

result :{
    uid       : 
    img_path  :
    text      : 
    res       :

    word_path :
    text_path : 
    pdf_path  :

}

'''

ALL_TYPES = ['.png', '.PNG', '.jpg', '.JPG', '.jpeg', '.JPEG']


### PDF

def create_pdf(input_, output="disk_report.pdf"):
    pdfmetrics.registerFont(TTFont("SimSun", "static/fonts/SimSun.ttf"))
    c = canvas.Canvas(output, (6 * inch, 10 * inch))
    c.setFont("SimSun", 12)

    def wrap(self, availWidth, availHeight):
        # work out widths array for breaking
        self.width = availWidth
        leftIndent = self.style.leftIndent
        first_line_width = availWidth - (leftIndent + self.style.firstLineIndent) - self.style.rightIndent
        later_widths = availWidth - leftIndent - self.style.rightIndent
        try:
            self.blPara = self.breakLinesCJK([first_line_width, later_widths])
        except:
            self.blPara = self.breakLines([first_line_width, later_widths])
        self.height = len(self.blPara.lines) * self.style.leading
        return (self.width, self.height)

    Paragraph.wrap = wrap
    styleSheet = getSampleStyleSheet()
    style = styleSheet['BodyText']
    style.fontName = "SimSun"
    style.fontSize = 16
    # 设置行距
    style.leading = 20
    # 首行缩进
    style.firstLineIndent = 32
    print(input_)
    Pa = Paragraph(input_, style)
    Pa.wrapOn(c, 6 * inch, 8 * inch)
    Pa.drawOn(c, 0, 5 * inch)

    c.showPage()
    c.save()


### session
SESSION = {}


class Session:
    def __init__(self, handler):
        self.handler = handler

    @staticmethod
    def _random_str():
        '''用随机数来作为session_id'''
        return str(uuid.uuid4())

    def _get_cookie_sid(self):
        '''获取cookie中的session_id'''
        cookie_sid = self.handler.get_secure_cookie("__session", None)
        return str(cookie_sid, encoding="utf-8") if cookie_sid else None

    def __setitem__(self, key, value):
        cookie_sid = self._get_cookie_sid()
        if not cookie_sid:
            cookie_sid = self._random_str()
        self.handler.set_secure_cookie("__session", cookie_sid)
        SESSION.setdefault(cookie_sid, {}).__setitem__(key, value)

    def __getitem__(self, key):
        cookie_sid = self._get_cookie_sid()
        content = SESSION.get(cookie_sid, None)
        return content.get(key, None) if content else None

    def __delitem__(self, key):
        cookie_sid = self._get_cookie_sid()
        if not cookie_sid or not SESSION.get(cookie_sid):
            raise KeyError(key)
        del SESSION[cookie_sid][key]


class BaseHandler(RequestHandler):

    def initialize(self):
        self.session = Session(self)


####### 接口

class LogoutHandler(BaseHandler):
    def get(self):
        self.session['name'] = None
        self.session['uid'] = None

        return self.redirect('/login')


class LoginHandler(BaseHandler):
    def get(self):
        return self.render('web/login.html')

    def post(self):
        user = self.get_body_argument('username', '')
        password = self.get_body_argument('password', '')
        mycol = db["users"]
        user_ = mycol.find_one({"name": user, "password": password},
                               {"_id": 1, "name": 1, "password": 1, " ocr_times": 1})
        print(user_)
        if user_ is None:
            return self.render("web/login_error.html", error="用户不存在")
        self.session['name'] = user_['name']
        self.session['uid'] = user_['_id']
        return self.redirect("/2")


class RegisterHandler(BaseHandler):
    def get(self):
        return self.render('web/register.html')

    def post(self):
        user = self.get_body_argument('username', '')
        password = self.get_body_argument('password', '')
        if user == '' or password == '':
            return self.render("web/register_error.html", error="用户名 密码不能为空")
        mycol = db["users"]
        old_user = mycol.find_one({"name": user})
        if old_user:
            return self.render("web/register_error.html", error="用户已存在")

        mycol = db["users"]
        new_user = {"name": user, "password": password, "ocr_times": 0}
        x = mycol.insert_one(new_user)
        if x:
            return self.redirect("/login")
        else:
            return self.render("web/register_error.html", error="注册失败，请冲重新注册")


class Main1Handler(BaseHandler):
    def get(self):
        if self.session['uid'] == None:  # 未登录
            return self.render('web/login_error.html', error="你还未登录")

        mycol = db["users"]
        user = mycol.find_one({"_id": self.session['uid']})
        print(user)
        if user == None:
            return self.render('web/login_error.html', error="用户不存在，请重新登录")
        return self.render('web/index1.html', ocr_times=user['ocr_times'], name=user['name'])


class Main2Handler(BaseHandler):
    def get(self):
        if self.session['uid'] == None:  # 未登录
            return self.render('web/login_error.html', error="你还未登录")

        mycol = db["users"]
        user = mycol.find_one({"_id": self.session['uid']})
        print(user)
        if user == None:
            return self.render('web/login_error.html', error="用户不存在，请重新登录")
        return self.render('web/index2.html', ocr_times=user['ocr_times'], name=user['name'])


class UploadFile2Handler(BaseHandler):

    def post(self):

        # 登录检查
        if self.session['uid'] == None:  # 未登录
            return self.render('web/login_error.html', error="你还未登录")

        mycol = db["users"]
        user = mycol.find_one({"_id": self.session['uid']})
        print(user)
        if user == None:
            return self.render('web/login_error.html', error="用户不存在，请重新登录")

        print(self.session)
        from settings import upload_path
        from ocr import OCR

        now_time = time.strftime('%Y-%m-%dT%H-%M-%S', time.localtime(time.time()))
        dir_prefix = now_time

        try:
            file_metas = self.request.files['file']
        except:
            return self.render('web/result.html', result='negative', result_header='失败', result_content='请选择需要识别的文件',
                               img_url="", rec_id="")

        for meta in file_metas:
            filename = meta['filename']
            # valid filetype
            if not os.path.splitext(filename)[1] in ALL_TYPES:
                return self.render('web/result.html', result='negative', result_header='上传失败',
                                   result_content='文件格式不支持', img_url="", rec_id="")
            try:
                os.makedirs(os.path.join(upload_path, dir_prefix))
            except:
                pass

            # save IMG file
            filepath = os.path.join(upload_path, dir_prefix, "img" + os.path.splitext(filename)[1])
            with open(filepath, 'wb') as up:
                up.write(meta['body'])

            # orc
            ocrinstance = OCR()
            res = ocrinstance.getResult(filepath)
            statusCode = res['code']
            status = '成功' if (statusCode == 1) else '失败'
            text = res['text']

            # save res
            respath = os.path.join(upload_path, dir_prefix, 'result.txt')
            with open(respath, 'w', encoding="utf-8") as info:
                info.write(status + '\n' + text)

            # save word
            docpath = os.path.join(upload_path, dir_prefix, 'result.docx')
            Doc = Document();
            Doc.add_heading("识别结果:" + status)
            Doc.add_paragraph(text)
            Doc.save(docpath)

            # save pdf
            pdfpath = os.path.join(upload_path, dir_prefix, 'result.pdf')
            c = canvas.Canvas(pdfpath)
            c.drawString(300, 1000, text)
            c.save()

            # 记录次数
            filepath = filepath.replace("/", "\\")
            print(filepath)
            relate_path = filepath.split('\\static\\')[1]
            print(relate_path)
            query = {"_id": self.session['uid']}
            newvalues = {"$inc": {"ocr_times": 1}}
            x = mycol.update_many(query, newvalues)

            res_col = db["results"]
            new_res = {"uid": user['_id'], 'img_path': relate_path, "text": text, "res": status, "doc_path": docpath,
                       "text_path": respath, "pdf_path": pdfpath}
            x2 = res_col.insert_one(new_res)
            print(x2)

            # render
            if (statusCode == 1):
                return self.render('web/result.html', result='positive', result_header=status, result_content=text,
                                   img_url=relate_path, rec_id=x2.inserted_id)
            else:
                return self.render('web/result.html', result='negative', result_header=status, result_content=text,
                                   img_url="", rec_id="")


class UploadFile1Handler(BaseHandler):

    def post(self):

        # 登录检查
        if self.session['uid'] == None:  # 未登录
            return self.finish({"code": -111, "msg": "unlogined ", "data": []})

        mycol = db["users"]
        user = mycol.find_one({"_id": self.session['uid']})
        print(user)
        if user == None:
            return self.finish({"code": -112, "msg": "user not exists", "data": []})

        from settings import upload_path
        from ocr import OCR

        now_time = time.strftime('%Y-%m-%dT%H-%M-%S', time.localtime(time.time()))
        dir_prefix = now_time

        try:
            print(111)
            base64ImgData = self.request.arguments['data'][0].decode("utf-8")
            # 去除base64图片前面的说明str
            base64ImgData = base64ImgData[base64ImgData.find(',') + 1:]
            imgData = base64.b64decode(base64ImgData)
        except:
            print(222)
            self.finish({"code": -113, "msg": "image encode error  ", "data": []})

        try:
            os.makedirs(os.path.join(upload_path, dir_prefix))
        except:
            print(333)
            self.finish({"code": -114, "msg": "mkdir error ", "data": []})

        # save file
        filepath = os.path.join(upload_path, dir_prefix, "img.png")
        with open(filepath, 'wb') as up:
            up.write(imgData)
        # orc
        ocrinstance = OCR()
        res = ocrinstance.getResult(filepath)
        statusCode = res['code']
        status = '成功' if (statusCode == 1) else '失败'
        text = res['text']

        # 记录次数
        filepath = filepath.replace("/", "\\")
        print(filepath)
        relate_path = filepath.split('\\static\\')[1]
        print(relate_path)
        query = {"_id": self.session['uid']}
        newvalues = {"$inc": {"ocr_times": 1}}
        x = mycol.update_many(query, newvalues)

        res_col = db["results"]
        new_res = {"uid": user['_id'], 'img_path': relate_path, "text": text, "res": status}
        x2 = res_col.insert_one(new_res)
        print(x2)

        # save res
        respath = os.path.join(upload_path, dir_prefix, 'result.txt')
        with open(respath, 'w', encoding="utf-8") as info:
            info.write(status + '\n' + text)

        # render
        print(444)
        if (statusCode == 1):
            print(555)
            self.finish({"code": 0, "msg": "succ", "data": str(ObjectId(x2.inserted_id))})
        else:
            print(666)
            self.finish({"code": -1, "msg": "failed ", "data": []})


class ResListHandler(BaseHandler):
    def get(self):
        # 登录检查
        if self.session['uid'] == None:  # 未登录
            return self.render('web/login_error.html', error="你还未登录")

        mycol = db["users"]
        user = mycol.find_one({"_id": self.session['uid']})
        print(user)
        if user == None:
            return self.render('web/login_error.html', error="用户不存在，请重新登录")

        res_col = db["results"]
        resArr = res_col.find({"uid": self.session['uid']})
        imgs = []
        for r in resArr:
            imgs.append(r)
        print('111111      ', imgs)

        return self.render('web/res_list.html', name=user['name'], ocr_times=user["ocr_times"], res_list=imgs)


class UpdateResHandler(BaseHandler):

    def post(self):

        try:
            # 登录检查
            if self.session['uid'] == None:  # 未登录
                return self.finish({"code": -5, "msg": "text is null", "data": []})

            mycol = db["users"]
            user = mycol.find_one({"_id": self.session['uid']})
            print(user)
            if user == None:
                return self.finish({code: -3, msg: 'not login', data: []})

            res_id = ObjectId(self.get_body_argument('res_id', ''))
            text = self.get_body_argument('text')
            print('res_id  ,  ', res_id)
            if res_id == "":
                return self.finish({"code": -1, "msg": "res_id is null", "data": []})

            if text == None:
                return self.finish({"code": -2, "msg": "text is null", "data": []})

            res_col = db["results"]
            query = {'_id': res_id}
            res_old = res_col.find_one(query)
            if res_old == None:
                return self.finish({"code": -6, "msg": 'result not exists', "data": []})

            newvalues = {"$set": {"text": text}}
            res = res_col.update_one(query, newvalues)
            print(res)
            return self.finish({"code": 0, "msg": "update succ", "data": []})
        except:
            return self.finish({"code": -8, "msg": "exception error ", "data": []})


class DownloadHandler(BaseHandler):
    def post(self):
        download_type = self.get_argument('type')
        res_id = self.get_argument('res_id')
        res_id_obj = ObjectId(res_id)

        if res_id == "":
            print(111)
            return self.finish({"code": -11, "msg": "res_id is null", "data": []})

        res_col = db["results"]
        query = {'_id': res_id_obj}
        res = res_col.find_one(query)
        if res == None:
            print(222)
            return self.finish({"code": -16, "msg": 'result not exists', "data": []})

        if download_type == "" or download_type == None:
            download_type = "word"
        suffix = ""

        if download_type == "word":
            suffix = '.docx'
        elif download_type == 'pdf':
            suffix = '.pdf'
        elif download_type == 'txt':
            suffix = '.txt'

        filepath = 'static/tmp/' + res_id + suffix

        if download_type == 'word':
            Doc = Document();
            Doc.add_heading("识别结果")
            Doc.add_paragraph(res['text'])
            Doc.save(filepath)
        elif download_type == 'pdf':
            # save pdf
            # pdfmetrics.registerFont(TTFont("SimSun", "static/fonts/SimSun.ttf"))
            # c = canvas.Canvas(filepath)
            # c.setFont("SimSun", 12)
            # c.drawString(textobject)
            # c.save()

            create_pdf(input_=res['text'], output=filepath)
        else:
            # save txt
            with open(filepath, 'w', encoding="utf-8") as info:
                info.write(res['text'])

        print(333)

        self.set_header('Content-Type', 'application/octet-stream')
        self.set_header('Content-Disposition', 'attachment; filename=' + filepath)
        with open(os.path.join('', filepath), 'rb') as f:
            while (True):
                data = f.read(200)
                if not data:
                    break
                self.write(data)
        print(444)
        self.finish()


class ResultHandler(BaseHandler):
    def get(self):

        if self.session['uid'] == None:  # 未登录
            return self.finish({"code": -5, "msg": "text is null", "data": []})

        mycol = db["users"]
        user = mycol.find_one({"_id": self.session['uid']})
        if user == None:
            return self.finish({code: -3, msg: 'not login', data: []})

        res_id = ObjectId(self.get_query_argument('res_id', ''))
        if res_id == "":
            print(111)
            return self.finish({"code": -11, "msg": "res_id is null", "data": []})

        res_col = db["results"]
        query = {'_id': res_id}
        res = res_col.find_one(query)
        if res == None:
            print(222)
            return self.finish({"code": -16, "msg": 'result not exists', "data": []})

        self.render('web/result.html', result='positive', result_header='成功', result_content=res['text'],
                    img_url=res['img_path'], rec_id=res['_id'])


if __name__ == "__main__":
    application = tornado.web.Application([
        (r"/", tornado.web.RedirectHandler, {"url": "/2", "permanent": False}),
        (r"/2", Main2Handler),
        (r"/1", Main1Handler),
        (r"/1/upload", UploadFile1Handler),
        (r"/2/upload", UploadFile2Handler),
        (r"/login", LoginHandler),
        (r"/logout", LogoutHandler),
        (r"/register", RegisterHandler),
        (r"/res_list", ResListHandler),
        (r"/update_res", UpdateResHandler),
        (r"/download_res", DownloadHandler),
        (r"/result", ResultHandler),

    ],
        static_path=static_path, cookie_secret="uYh7EQnp2XdTP1o/Vo=")

    application.listen(PORT, address=IP)
    tornado.ioloop.IOLoop.instance().start()

